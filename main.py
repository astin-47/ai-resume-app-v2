# =============================================================
# main.py  —  AI Resume Builder  v5.1  (Auth + Limits + Razorpay)
# =============================================================
#
#  FREE plan  : 5 lifetime generations (resume + cover letter)
#               ATS checks are always free
#  PRO plan   : unlimited everything  —  ₹499/month
#
#  Razorpay payment flow:
#  ──────────────────────
#  1. User clicks "Upgrade" in the frontend
#  2. Frontend calls POST /create-order
#  3. Backend calls Razorpay API → gets back an order_id
#  4. Frontend opens the Razorpay checkout POPUP (user stays on page)
#  5. User pays via UPI / card / netbanking / wallet
#  6. Razorpay returns payment_id + signature to the frontend
#  7. Frontend calls POST /verify-payment with those details
#  8. Backend verifies the HMAC signature (proves payment is real)
#  9. Backend upgrades user to PRO in Supabase
#
#  Why verify on the backend?
#  ──────────────────────────
#  Never trust the frontend alone for payments. Anyone could call
#  /verify-payment with fake data. The HMAC check uses your
#  Razorpay secret key (which only your server knows) to confirm
#  the payment_id is genuine and was not tampered with.
#
#  ENDPOINTS:
#  ──────────
#   GET  /               — health check
#   GET  /my-usage       — user's plan + usage count (auth required)
#   POST /upload         — extract text (auth required)
#   POST /generate       — AI generation (auth + plan check)
#   POST /create-order   — create Razorpay order (auth required)
#   POST /verify-payment — verify & activate Pro (auth required)
#
# =============================================================

import os, io, json, tempfile, shutil, subprocess, hmac, hashlib
from fastapi import FastAPI, File, UploadFile, Form, HTTPException, Request
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import JSONResponse, FileResponse

from dotenv import load_dotenv
load_dotenv()

import PyPDF2
import httpx
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from openai import OpenAI

# ── Clients & config ──────────────────────────────────────────
client               = OpenAI(api_key=os.getenv("OPENAI_API_KEY"))
SUPABASE_URL         = os.getenv("SUPABASE_URL")
SUPABASE_ANON_KEY    = os.getenv("SUPABASE_ANON_KEY")
SUPABASE_SERVICE_KEY = os.getenv("SUPABASE_SERVICE_KEY")

# Razorpay credentials — set in Render environment variables
RAZORPAY_KEY_ID     = os.getenv("RAZORPAY_KEY_ID")      # starts with rzp_live_
RAZORPAY_KEY_SECRET = os.getenv("RAZORPAY_KEY_SECRET")  # keep this secret, server only

# Plan config
FREE_LIMIT    = 5          # free generations before paywall
PRO_AMOUNT    = 49900      # amount in paise (₹499 = 49900 paise)
PRO_CURRENCY  = "INR"
PRO_PLAN_NAME = "ResumeForge Pro"

ALLOWED_EXT = {".pdf", ".docx", ".txt"}
DOCX_MIME   = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# ── App ───────────────────────────────────────────────────────
app = FastAPI(title="AI Resume Builder", version="5.1.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=False,
    allow_methods=["*"],
    allow_headers=["*"],
)


# =============================================================
#  AUTH — verify Supabase JWT, return user_id
# =============================================================
async def get_current_user(request: Request) -> str:
    auth = request.headers.get("Authorization", "")
    if not auth.startswith("Bearer "):
        raise HTTPException(401, "Not logged in. Please sign in.")
    token = auth.split(" ", 1)[1]
    try:
        async with httpx.AsyncClient() as h:
            r = await h.get(
                f"{SUPABASE_URL}/auth/v1/user",
                headers={"Authorization": f"Bearer {token}",
                         "apikey": SUPABASE_ANON_KEY},
                timeout=10,
            )
    except Exception:
        raise HTTPException(503, "Auth service unreachable. Try again.")
    if r.status_code != 200:
        raise HTTPException(401, "Session expired. Please log in again.")
    return r.json().get("id", "unknown")


# =============================================================
#  SUPABASE PLAN HELPERS
# =============================================================
def _svc():
    """Service-role headers that bypass Row Level Security."""
    return {
        "apikey":        SUPABASE_SERVICE_KEY,
        "Authorization": f"Bearer {SUPABASE_SERVICE_KEY}",
        "Content-Type":  "application/json",
    }


async def get_user_plan(user_id: str) -> dict:
    """
    Returns {"plan": "free"|"pro", "usage_count": int}.
    Creates a free plan row automatically if none exists.
    """
    async with httpx.AsyncClient() as h:
        r = await h.get(
            f"{SUPABASE_URL}/rest/v1/user_plans",
            headers=_svc(),
            params={"user_id": f"eq.{user_id}", "select": "plan,usage_count"},
            timeout=10,
        )
    rows = r.json()
    if rows:
        return rows[0]
    # First time — create a free plan row
    async with httpx.AsyncClient() as h:
        await h.post(
            f"{SUPABASE_URL}/rest/v1/user_plans",
            headers={**_svc(), "Prefer": "return=minimal"},
            json={"user_id": user_id, "plan": "free", "usage_count": 0},
            timeout=10,
        )
    return {"plan": "free", "usage_count": 0}


async def increment_usage(user_id: str):
    """Add 1 to usage_count atomically via a Supabase RPC function."""
    async with httpx.AsyncClient() as h:
        await h.post(
            f"{SUPABASE_URL}/rest/v1/rpc/increment_usage",
            headers=_svc(),
            json={"uid": user_id},
            timeout=10,
        )


async def upgrade_to_pro(user_id: str, razorpay_payment_id: str):
    """
    Set user's plan to 'pro' after payment is verified.
    Stores the Razorpay payment ID for your records.
    """
    async with httpx.AsyncClient() as h:
        await h.patch(
            f"{SUPABASE_URL}/rest/v1/user_plans",
            headers={**_svc(), "Prefer": "return=minimal"},
            params={"user_id": f"eq.{user_id}"},
            json={"plan": "pro", "razorpay_payment_id": razorpay_payment_id},
            timeout=10,
        )


# =============================================================
#  RAZORPAY HELPERS
# =============================================================
async def razorpay_create_order(amount: int, currency: str,
                                 receipt: str, notes: dict) -> dict:
    """
    Calls the Razorpay Orders API to create a new order.
    Returns the full order object including the order_id.

    We use HTTP Basic Auth:
      username = RAZORPAY_KEY_ID
      password = RAZORPAY_KEY_SECRET
    """
    async with httpx.AsyncClient() as h:
        r = await h.post(
            "https://api.razorpay.com/v1/orders",
            auth=(RAZORPAY_KEY_ID, RAZORPAY_KEY_SECRET),
            json={
                "amount":   amount,    # in paise (smallest currency unit)
                "currency": currency,
                "receipt":  receipt,   # your internal reference
                "notes":    notes,     # any metadata you want to store
            },
            timeout=15,
        )
    if r.status_code != 200:
        raise HTTPException(500, f"Razorpay order creation failed: {r.text}")
    return r.json()


def verify_razorpay_signature(order_id: str, payment_id: str,
                               signature: str) -> bool:
    """
    Verifies the HMAC-SHA256 signature that Razorpay sends after payment.

    How it works:
      1. Razorpay concatenates order_id + "|" + payment_id
      2. Signs it with your KEY_SECRET using HMAC-SHA256
      3. Sends you the signature
      4. You do the same calculation and compare
      5. If they match → payment is genuine
      6. If they don't → someone tampered with the data

    This is the ONLY reliable way to confirm a payment is real.
    Never skip this check.
    """
    message  = f"{order_id}|{payment_id}"
    expected = hmac.new(
        RAZORPAY_KEY_SECRET.encode("utf-8"),
        message.encode("utf-8"),
        hashlib.sha256,
    ).hexdigest()
    # Use hmac.compare_digest to prevent timing attacks
    return hmac.compare_digest(expected, signature)


# =============================================================
#  TEXT EXTRACTION
# =============================================================
def extract_text(data: bytes, filename: str) -> str:
    ext = os.path.splitext(filename)[1].lower()
    if ext not in ALLOWED_EXT:
        raise ValueError(f"'{ext}' not supported. Upload PDF, DOCX, or TXT.")
    text = ""
    if ext == ".pdf":
        try:
            reader = PyPDF2.PdfReader(io.BytesIO(data))
            for page in reader.pages:
                chunk = page.extract_text()
                if chunk: text += chunk + "\n"
        except Exception as e:
            raise ValueError(f"Could not read PDF: {e}")
    elif ext == ".docx":
        try:
            doc = Document(io.BytesIO(data))
            for para in doc.paragraphs:
                text += para.text + "\n"
        except Exception as e:
            raise ValueError(f"Could not read DOCX: {e}")
    elif ext == ".txt":
        text = data.decode("utf-8", errors="replace")
    text = text.strip()
    if not text:
        raise ValueError("File is empty or has no extractable text.")
    return text


# =============================================================
#  OPENAI HELPERS
# =============================================================
def ask_ai(system: str, user: str, max_tokens: int = 3500) -> str:
    resp = client.chat.completions.create(
        model="gpt-4o-mini",
        messages=[{"role":"system","content":system},
                  {"role":"user",  "content":user}],
        temperature=0.65,
        max_tokens=max_tokens,
    )
    return resp.choices[0].message.content.strip()


def parse_json(raw: str) -> dict:
    clean = raw.strip()
    if clean.startswith("```"):
        parts = clean.split("```")
        clean = parts[1]
        if clean.lower().startswith("json"):
            clean = clean[4:]
    return json.loads(clean.strip())


def run_ats_check(resume_text: str, job_description: str) -> dict:
    system = """You are an expert ATS analyst.
Analyse the resume against the job description.
Return ONLY valid JSON — no markdown, no explanation:
{
  "score": <integer 0-100>,
  "level": "<Excellent|Good|Fair|Poor>",
  "matched_keywords": ["keyword in resume"],
  "missing_keywords": ["keyword missing from resume"],
  "suggestions": ["specific actionable improvement"]
}
Scoring: 90-100 Excellent, 70-89 Good, 50-69 Fair, 0-49 Poor"""
    user = (f"Resume:\n{resume_text}\n\nJob Description:\n"
            f"{job_description or 'No job description — score on general best practices.'}")
    try:
        return parse_json(ask_ai(system, user, max_tokens=1500))
    except Exception:
        return {"score":50,"level":"Fair","matched_keywords":[],
                "missing_keywords":[],"suggestions":["Could not parse ATS analysis."]}


def get_optimised_resume_data(resume_text: str, job_description: str, ats: dict) -> dict:
    missing = ", ".join(ats.get("missing_keywords",[])) or "none"
    matched = ", ".join(ats.get("matched_keywords",[])) or "none"
    score   = ats.get("score", 50)
    sugg    = "\n".join(f"- {s}" for s in ats.get("suggestions",[]))
    system  = f"""You are an expert resume writer and ATS optimisation specialist.
Current ATS score: {score}/100. Matched: {matched}. Missing: {missing}.
Suggestions:\n{sugg}

STRICT RULES:
1. Do NOT invent jobs, companies, dates, degrees, or qualifications.
2. Rephrase bullets to include missing keywords where genuinely applicable.
3. Add missing keywords to Skills only if legitimately related.
4. Rewrite the Professional Summary to align with the target role.
5. Most relevant bullets first. Return [] for empty sections.

Return ONLY valid JSON (no markdown):
{{"name":"","phone":"","email":"","linkedin":"","summary":"",
"skills":[{{"category":"","items":""}}],
"experience":[{{"title":"","company":"","dates":"","bullets":[]}}],
"education":[{{"institution":"","degree":"","dates":"","grade":""}}],
"certifications":[{{"name":"","year":""}}],
"projects":[{{"name":"","tech":"","bullets":[]}}]}}"""
    raw = ask_ai(system,
                 f"Original Resume:\n{resume_text}\n\nTarget Job:\n{job_description}",
                 max_tokens=4000)
    try:
        return parse_json(raw)
    except json.JSONDecodeError:
        raise ValueError("AI returned unexpected data. Please try again.")


def get_optimised_cover_letter(resume_text: str, job_description: str, ats: dict) -> str:
    missing = ", ".join(ats.get("missing_keywords",[])) or "none"
    matched = ", ".join(ats.get("matched_keywords",[])) or "none"
    system  = f"""You are a world-class cover letter writer.
ATS score: {ats.get('score',50)}/100. Matched: {matched}. Missing: {missing}.
Write a cover letter: strong hook, 2 body paragraphs using real experience,
strong closing. Max 380 words. Use the candidate's real name."""
    return ask_ai(system,
                  f"Resume:\n{resume_text}\n\nJob Description:\n{job_description}",
                  max_tokens=1500)


# =============================================================
#  DOCX BUILDER
# =============================================================
def _font(run, size=11, bold=False, italic=False):
    run.font.name=  "Calibri"
    run.font.size=  Pt(size)
    run.font.bold=  bold
    run.font.italic=italic

def _bottom_border(p):
    pBdr=OxmlElement("w:pBdr"); b=OxmlElement("w:bottom")
    b.set(qn("w:val"),"single"); b.set(qn("w:sz"),"6")
    b.set(qn("w:space"),"1"); b.set(qn("w:color"),"000000")
    pBdr.append(b); p._p.get_or_add_pPr().append(pBdr)

def _section(doc, title):
    p=doc.add_paragraph()
    p.paragraph_format.space_before=Pt(10)
    p.paragraph_format.space_after=Pt(2)
    _font(p.add_run(title.upper()),size=13,bold=True)
    _bottom_border(p)

def _right_tab(p):
    pPr=p._p.get_or_add_pPr(); tabs=OxmlElement("w:tabs")
    tab=OxmlElement("w:tab"); tab.set(qn("w:val"),"right")
    tab.set(qn("w:pos"),"9360"); tabs.append(tab); pPr.append(tabs)

def _base_doc():
    doc=Document()
    for s in doc.sections:
        s.top_margin=s.bottom_margin=s.left_margin=s.right_margin=Inches(0.5)
    n=doc.styles["Normal"]; n.font.name="Calibri"
    n.font.size=Pt(11); n.paragraph_format.line_spacing=Pt(13)
    return doc

def build_resume_docx(d: dict) -> bytes:
    doc=_base_doc()
    p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(2)
    _font(p.add_run(d.get("name","")),size=22,bold=True)
    parts=[x for x in [d.get("phone"),d.get("email"),d.get("linkedin")] if x and x.strip()]
    if parts:
        cp=doc.add_paragraph(); cp.paragraph_format.space_after=Pt(6)
        _font(cp.add_run("  |  ".join(parts)),size=10,italic=True)
    if d.get("summary","").strip():
        _section(doc,"Professional Summary")
        p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.15)
        p.paragraph_format.space_after=Pt(4); _font(p.add_run(d["summary"]))
    skills=[s for s in d.get("skills",[]) if s.get("category","").strip()]
    if skills:
        _section(doc,"Technical Skills")
        for sk in skills:
            p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.15)
            p.paragraph_format.space_after=Pt(1)
            _font(p.add_run(sk.get("category","")+": "),bold=True)
            _font(p.add_run(sk.get("items","")))
    exp=[j for j in d.get("experience",[]) if j.get("title","").strip()]
    if exp:
        _section(doc,"Experience")
        for job in exp:
            p=doc.add_paragraph()
            p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(1)
            _font(p.add_run(job.get("title","")),bold=True)
            _font(p.add_run("  |  "+job.get("company","")))
            p.add_run("\t"); _font(p.add_run(job.get("dates","")),size=10,italic=True)
            _right_tab(p)
            for b in job.get("bullets",[]):
                if not b.strip(): continue
                bp=doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.left_indent=Inches(0.3)
                bp.paragraph_format.space_after=Pt(1); _font(bp.add_run(b))
    proj=[p for p in d.get("projects",[]) if p.get("name","").strip()]
    if proj:
        _section(doc,"Projects")
        for pr in proj:
            p=doc.add_paragraph()
            p.paragraph_format.space_before=Pt(6); p.paragraph_format.space_after=Pt(1)
            _font(p.add_run(pr.get("name","")),bold=True)
            if pr.get("tech","").strip(): _font(p.add_run("  |  "+pr["tech"]),italic=True)
            for b in pr.get("bullets",[]):
                if not b.strip(): continue
                bp=doc.add_paragraph(style="List Bullet")
                bp.paragraph_format.left_indent=Inches(0.3)
                bp.paragraph_format.space_after=Pt(1); _font(bp.add_run(b))
    certs=[c for c in d.get("certifications",[]) if c.get("name","").strip()]
    if certs:
        _section(doc,"Awards and Certifications")
        for c in certs:
            bp=doc.add_paragraph(style="List Bullet")
            bp.paragraph_format.left_indent=Inches(0.3); bp.paragraph_format.space_after=Pt(1)
            _font(bp.add_run(c.get("name","")),bold=True)
            if c.get("year","").strip(): _font(bp.add_run("  |  "+c["year"]),italic=True)
    edu=[e for e in d.get("education",[]) if e.get("institution","").strip()]
    if edu:
        _section(doc,"Education")
        for e in edu:
            p=doc.add_paragraph(); p.paragraph_format.left_indent=Inches(0.15)
            p.paragraph_format.space_after=Pt(2)
            _font(p.add_run(e.get("institution","")),bold=True)
            if e.get("grade","").strip(): _font(p.add_run("  "+e["grade"]),bold=True)
            _font(p.add_run("  •  "+e.get("degree","")),italic=True)
            if e.get("dates","").strip(): _font(p.add_run("  •  "+e["dates"]),size=10,italic=True)
    buf=io.BytesIO(); doc.save(buf); buf.seek(0); return buf.read()


def build_cover_letter_docx(cover_text: str) -> bytes:
    doc=Document()
    for s in doc.sections:
        s.top_margin=s.bottom_margin=s.left_margin=s.right_margin=Inches(1.0)
    for line in cover_text.strip().split("\n"):
        line=line.strip()
        if not line: doc.add_paragraph(); continue
        p=doc.add_paragraph(); p.paragraph_format.space_after=Pt(8); _font(p.add_run(line))
    buf=io.BytesIO(); doc.save(buf); buf.seek(0); return buf.read()


def docx_to_pdf(docx_bytes: bytes):
    with tempfile.TemporaryDirectory() as tmp:
        dp=os.path.join(tmp,"resume.docx"); pp=os.path.join(tmp,"resume.pdf")
        with open(dp,"wb") as f: f.write(docx_bytes)
        lo=shutil.which("soffice") or shutil.which("libreoffice")
        if lo:
            try:
                r=subprocess.run([lo,"--headless","--convert-to","pdf","--outdir",tmp,dp],
                                 capture_output=True,timeout=60)
                if r.returncode==0 and os.path.exists(pp):
                    return open(pp,"rb").read()
            except Exception: pass
        try:
            from docx2pdf import convert
            convert(dp,pp)
            if os.path.exists(pp): return open(pp,"rb").read()
        except Exception: pass
    return None


def temp_response(data, suffix, media, filename, extra_headers=None):
    tmp=tempfile.NamedTemporaryFile(delete=False,suffix=suffix)
    tmp.write(data); tmp.close()
    return FileResponse(path=tmp.name,media_type=media,
                        filename=filename,headers=extra_headers or {})


# =============================================================
#  ENDPOINT  GET /
# =============================================================
@app.get("/")
def root():
    return {"status":"ok","message":"AI Resume Builder v5.1 (Razorpay) is running ✅"}


# =============================================================
#  ENDPOINT  GET /my-usage
# =============================================================
@app.get("/my-usage")
async def my_usage(request: Request):
    user_id   = await get_current_user(request)
    plan      = await get_user_plan(user_id)
    remaining = max(0, FREE_LIMIT - plan["usage_count"]) if plan["plan"] == "free" else None
    return {
        "plan":        plan["plan"],
        "usage_count": plan["usage_count"],
        "free_limit":  FREE_LIMIT,
        "remaining":   remaining,
    }


# =============================================================
#  ENDPOINT  POST /upload
# =============================================================
@app.post("/upload")
async def upload(request: Request, file: UploadFile = File(...)):
    await get_current_user(request)
    if not file or not file.filename:
        raise HTTPException(400, "No file received.")
    data = await file.read()
    if not data: raise HTTPException(400, "File is empty.")
    try:
        text = extract_text(data, file.filename)
    except ValueError as e:
        raise HTTPException(422, str(e))
    return {"filename": file.filename, "characters": len(text), "text": text}


# =============================================================
#  ENDPOINT  POST /generate
# =============================================================
@app.post("/generate")
async def generate(
    request:         Request,
    file:            UploadFile = File(...),
    job_description: str        = Form(default=""),
    option:          str        = Form(default="resume"),
    output_format:   str        = Form(default="docx"),
):
    user_id = await get_current_user(request)

    # Plan check — ATS is always free, resume/cover count against limit
    if option in {"resume", "cover"}:
        plan = await get_user_plan(user_id)
        if plan["plan"] == "free" and plan["usage_count"] >= FREE_LIMIT:
            raise HTTPException(
                status_code=402,
                detail={
                    "code":    "limit_reached",
                    "message": f"You have used all {FREE_LIMIT} free generations.",
                    "used":    plan["usage_count"],
                    "limit":   FREE_LIMIT,
                }
            )

    if option not in {"resume","cover","ats"}:
        raise HTTPException(400, f"Unknown option '{option}'.")
    if output_format not in {"docx","pdf","txt"}:
        raise HTTPException(400, f"Unknown format '{output_format}'.")
    if not file or not file.filename:
        raise HTTPException(400, "Please attach your resume file.")

    raw_data = await file.read()
    if not raw_data: raise HTTPException(400, "Uploaded file is empty.")
    try:
        resume_text = extract_text(raw_data, file.filename)
    except ValueError as e:
        raise HTTPException(422, str(e))

    if option in {"resume","cover"}:
        await increment_usage(user_id)

    if option == "ats":
        try:
            return JSONResponse(content=run_ats_check(resume_text, job_description))
        except Exception as e:
            raise HTTPException(500, f"ATS check error: {e}")

    if option == "cover":
        if not job_description.strip():
            raise HTTPException(400, "A job description is required for a cover letter.")
        try:
            ats        = run_ats_check(resume_text, job_description)
            cover_text = get_optimised_cover_letter(resume_text, job_description, ats)
        except Exception as e:
            raise HTTPException(500, f"Error: {e}")
        if output_format == "txt":
            return temp_response(cover_text.encode(),".txt","text/plain","cover_letter.txt")
        docx_bytes = build_cover_letter_docx(cover_text)
        if output_format == "docx":
            return temp_response(docx_bytes,".docx",DOCX_MIME,"cover_letter.docx")
        pdf = docx_to_pdf(docx_bytes)
        if pdf: return temp_response(pdf,".pdf","application/pdf","cover_letter.pdf")
        return temp_response(docx_bytes,".docx",DOCX_MIME,"cover_letter.docx",
                             {"X-Fallback":"pdf-unavailable"})

    if option == "resume":
        try:
            ats  = run_ats_check(resume_text, job_description)
            data = get_optimised_resume_data(resume_text, job_description, ats)
        except ValueError as e:
            raise HTTPException(422, str(e))
        except Exception as e:
            raise HTTPException(500, f"Error: {e}")
        if output_format == "txt":
            txt = ask_ai("Rewrite as clean plain text. Headers ALL CAPS with dashes. "
                        "Bullet points start with '- '. No markdown.", resume_text)
            return temp_response(txt.encode(),".txt","text/plain","resume.txt")
        try:
            docx_bytes = build_resume_docx(data)
        except Exception as e:
            raise HTTPException(500, f"Error building document: {e}")
        if output_format == "docx":
            return temp_response(docx_bytes,".docx",DOCX_MIME,"resume.docx")
        pdf = docx_to_pdf(docx_bytes)
        if pdf: return temp_response(pdf,".pdf","application/pdf","resume.pdf")
        return temp_response(docx_bytes,".docx",DOCX_MIME,"resume.docx",
                             {"X-Fallback":"pdf-unavailable"})


# =============================================================
#  ENDPOINT  POST /create-order
#
#  Called by the frontend when the user clicks "Upgrade".
#  Creates a Razorpay order and returns the order_id + amount
#  so the frontend can open the Razorpay checkout popup.
# =============================================================
@app.post("/create-order")
async def create_order(request: Request):
    user_id = await get_current_user(request)

    try:
        order = await razorpay_create_order(
            amount   = PRO_AMOUNT,
            currency = PRO_CURRENCY,
            receipt  = f"receipt_{user_id[:8]}",   # short internal reference
            notes    = {
                "user_id":  user_id,
                "plan":     "pro",
                "product":  PRO_PLAN_NAME,
            },
        )
    except Exception as e:
        raise HTTPException(500, f"Could not create order: {e}")

    # Return what the frontend needs to open the Razorpay popup
    return {
        "order_id": order["id"],
        "amount":   order["amount"],     # in paise
        "currency": order["currency"],
        "key_id":   RAZORPAY_KEY_ID,     # public key — safe to send to frontend
    }


# =============================================================
#  ENDPOINT  POST /verify-payment
#
#  Called by the frontend AFTER the user pays successfully.
#  Razorpay gives the frontend 3 values:
#    razorpay_order_id   — the order we created
#    razorpay_payment_id — the unique payment reference
#    razorpay_signature  — HMAC proof the payment is genuine
#
#  We verify the signature, then upgrade the user to Pro.
# =============================================================
@app.post("/verify-payment")
async def verify_payment(request: Request):
    user_id = await get_current_user(request)
    body    = await request.json()

    order_id   = body.get("razorpay_order_id",   "")
    payment_id = body.get("razorpay_payment_id", "")
    signature  = body.get("razorpay_signature",  "")

    if not all([order_id, payment_id, signature]):
        raise HTTPException(400, "Missing payment details.")

    # Verify the HMAC signature — this is the security check
    if not verify_razorpay_signature(order_id, payment_id, signature):
        raise HTTPException(400, "Payment signature verification failed. "
                                 "Please contact support.")

    # Signature is valid — upgrade the user to Pro
    try:
        await upgrade_to_pro(user_id, payment_id)
    except Exception as e:
        raise HTTPException(500, f"Could not activate Pro: {e}")

    return {
        "success":    True,
        "plan":       "pro",
        "payment_id": payment_id,
        "message":    "Payment verified. Welcome to ResumeForge Pro! 🎉",
    }
